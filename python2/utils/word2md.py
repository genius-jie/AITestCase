#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word文档读取工具
用于读取Word文档内容并转换为Markdown格式
"""

import sys
from pathlib import Path
from docx import Document


def read_word_to_markdown(docx_path):
    """
    读取Word文档并转换为Markdown格式
    
    Args:
        docx_path: Word文档路径
        
    Returns:
        str: Markdown格式的内容
    """
    try:
        doc = Document(docx_path)
        markdown_content = []
        
        # 按照文档原始顺序遍历所有元素（段落和表格）
        for element in doc.element.body:
            # 判断元素类型
            if element.tag.endswith('p'):
                # 段落
                for para in doc.paragraphs:
                    if para._element == element:
                        text = para.text.strip()
                        if not text:
                            markdown_content.append("")
                            continue
                            
                        # 处理标题样式
                        style_name = para.style.name if para.style else "Normal"
                        if style_name.startswith('Heading'):
                            level = style_name.replace('Heading ', '')
                            if level.isdigit():
                                markdown_content.append(f"{'#' * int(level)} {text}")
                            else:
                                markdown_content.append(f"## {text}")
                        else:
                            markdown_content.append(text)
                        break
            elif element.tag.endswith('tbl'):
                # 表格
                for table in doc.tables:
                    if table._element == element:
                        markdown_content.append("")
                        for i, row in enumerate(table.rows):
                            row_data = [cell.text.strip() for cell in row.cells]
                            if i == 0:
                                markdown_content.append("| " + " | ".join(row_data) + " |")
                                markdown_content.append("| " + " | ".join(["---"] * len(row_data)) + " |")
                            else:
                                markdown_content.append("| " + " | ".join(row_data) + " |")
                        markdown_content.append("")
                        break
        
        return "\n".join(markdown_content)
    
    except Exception as e:
        import traceback
        return f"错误：无法读取文档 - {str(e)}\n详细错误：\n{traceback.format_exc()}"


def main():
    """主函数"""
    if len(sys.argv) < 2:
        print("使用方法: python read_word.py <word文档路径>")
        sys.exit(1)
    
    docx_path = sys.argv[1]
    
    if not Path(docx_path).exists():
        print(f"错误：文件不存在 - {docx_path}")
        sys.exit(1)
    
    markdown_content = read_word_to_markdown(docx_path)
    
    # 自动生成输出路径：将.docx替换为.md
    docx_path_obj = Path(docx_path)
    output_path = docx_path_obj.with_suffix('.md')
    
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(markdown_content)
    print(f"成功转换文档到: {output_path}")


if __name__ == "__main__":
    main()